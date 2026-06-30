#!/usr/bin/env python3
"""将 Markdown 系统说明书转换为 Word (.docx)，适配 GB/T 8567 风格排版。"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
DEFAULT_MD = ROOT / "docs" / "K-CARD系统说明书.md"
DEFAULT_DOCX = ROOT / "docs" / "K-CARD_System_Manual_v1.5.0.docx"

FONT_BODY = "宋体"
FONT_HEADING = "黑体"
FONT_CODE = "Consolas"


def set_run_font(run, name=FONT_BODY, size=12, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def setup_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = FONT_BODY
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    normal.font.size = Pt(12)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(6)

    for level, size in [(1, 16), (2, 14), (3, 13), (4, 12)]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = FONT_HEADING
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HEADING)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)


def add_cover(doc: Document, title: str, version: str, date: str):
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    set_run_font(r, FONT_HEADING, 22, bold=True, color=RGBColor(0x1A, 0x3A, 0x5C))

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("系统说明书")
    set_run_font(r2, FONT_HEADING, 18, bold=True)

    doc.add_paragraph()

    meta = [
        ("文档版本", version),
        ("编制日期", date),
        ("文档类型", "软件设计说明书 / 用户手册"),
        ("适用标准", "参照 GB/T 8567-2006《计算机软件文档编制规范》"),
        ("项目名称", "K-CARD KPOP 小卡交易平台"),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    table.style = "Table Grid"
    for i, (k, v) in enumerate(meta):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
        for c in table.rows[i].cells:
            for para in c.paragraphs:
                for run in para.runs:
                    set_run_font(run, FONT_BODY, 12)

    doc.add_page_break()


def parse_table_row(line: str):
    line = line.strip()
    if not line.startswith("|"):
        return None
    cells = [c.strip() for c in line.strip("|").split("|")]
    return cells


def is_separator_row(cells):
    return all(re.match(r"^:?-+:?$", c.replace(" ", "")) for c in cells if c)


def add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci in range(cols):
            text = row[ci] if ci < len(row) else ""
            cell = table.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            parts = re.split(r"(\*\*[^*]+\*\*)", text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    set_run_font(run, FONT_BODY, 11, bold=True)
                else:
                    run = p.add_run(part)
                    set_run_font(run, FONT_BODY, 11)
    doc.add_paragraph()


def add_rich_paragraph(doc: Document, text: str, style=None, indent=False):
    p = doc.add_paragraph(style=style)
    if indent:
        p.paragraph_format.left_indent = Cm(0.75)
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            set_run_font(run, FONT_BODY, 12, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            set_run_font(run, FONT_CODE, 10)
        else:
            run = p.add_run(part)
            set_run_font(run, FONT_BODY, 12)
    return p


def add_code_block(doc: Document, lines: list[str]):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    text = "\n".join(lines)
    run = p.add_run(text)
    set_run_font(run, FONT_CODE, 9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def convert_md_to_docx(md_path: Path, docx_path: Path):
    content = md_path.read_text(encoding="utf-8")
    lines = content.splitlines()

    version = "1.5.0"
    date = "2026-06-30"
    m_ver = re.search(r"版本[：:]\s*([\d.]+)", content)
    m_date = re.search(r"更新日期[：:]\s*([\d-]+)", content)
    if m_ver:
        version = m_ver.group(1)
    if m_date:
        date = m_date.group(1)

    doc = Document()
    setup_styles(doc)
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    add_cover(doc, "K-CARD", version, date)

    i = 0
    in_code = False
    code_buf: list[str] = []
    table_buf: list[list[str]] = []
    skip_until_chapter = True

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if skip_until_chapter:
            if stripped.startswith("## ") and not stripped.startswith("## 目录"):
                skip_until_chapter = False
            else:
                i += 1
                continue

        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_buf)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if stripped.startswith("|"):
            cells = parse_table_row(line)
            if cells is not None:
                if not is_separator_row(cells):
                    table_buf.append(cells)
                i += 1
                continue
        elif table_buf:
            add_table(doc, table_buf)
            table_buf = []

        if stripped.startswith("# ") and not stripped.startswith("## "):
            i += 1
            continue

        if stripped.startswith("## "):
            title = stripped[3:].strip()
            if title == "目录":
                i += 1
                continue
            doc.add_heading(title, level=1)
            i += 1
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
            i += 1
            continue

        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:].strip(), level=3)
            i += 1
            continue

        if stripped.startswith("> "):
            add_rich_paragraph(doc, stripped[2:], indent=True)
            i += 1
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            parts = re.split(r"(\*\*[^*]+\*\*)", stripped[2:])
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    set_run_font(run, FONT_BODY, 12, bold=True)
                elif part:
                    run = p.add_run(part)
                    set_run_font(run, FONT_BODY, 12)
            i += 1
            continue

        if re.match(r"^\d+\.\s", stripped):
            add_rich_paragraph(doc, stripped)
            i += 1
            continue

        if stripped:
            add_rich_paragraph(doc, stripped)

        i += 1

    if table_buf:
        add_table(doc, table_buf)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    print(f"Generated: {docx_path}")


if __name__ == "__main__":
    md = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_MD
    out = Path(sys.argv[2]) if len(sys.argv) > 2 else DEFAULT_DOCX
    convert_md_to_docx(md, out)
